"""
Essay formatter for MLA 9 and APA 7 student-paper defaults.
Uses OpenAI to detect structure and run citation/content checks.

Usage:
    python mla_formatter.py input.docx
    python mla_formatter.py input.docx --style apa
"""

import argparse, json, os, re, sys
from typing import Any
from dotenv import load_dotenv

# Load .env file from the script's directory
load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── AI Analysis ──────────────────────────────────────────────────────────────

OPENAI_MODEL = (os.getenv("OPENAI_MODEL", "gpt-5.2") or "gpt-5.2").strip()


def _env_float(name: str, default: float) -> float:
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    try:
        value = float(raw)
        return value if value > 0 else default
    except Exception:
        return default


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    try:
        value = int(raw)
        return value if value >= 0 else default
    except Exception:
        return default


def _build_openai_client(api_key: str):
    from openai import OpenAI

    timeout_seconds = _env_float("OPENAI_TIMEOUT_SECONDS", 45.0)
    max_retries = _env_int("OPENAI_MAX_RETRIES", 2)
    return OpenAI(api_key=api_key, timeout=timeout_seconds, max_retries=max_retries)


def _extract_message_text(resp: Any) -> str:
    try:
        content = resp.choices[0].message.content
    except Exception:
        return ""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        chunks = []
        for part in content:
            if isinstance(part, dict):
                text = part.get("text", "")
            else:
                text = getattr(part, "text", "")
            if text:
                chunks.append(text)
        return "\n".join(chunks)
    return str(content or "")


def _strip_json_fences(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return raw


def _chat_completion_json(client, prompt: str, schema_name: str, schema: dict):
    base_kwargs = {
        "model": OPENAI_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0,
    }
    try:
        resp = client.chat.completions.create(
            **base_kwargs,
            response_format={
                "type": "json_schema",
                "json_schema": {
                    "name": schema_name,
                    "strict": True,
                    "schema": schema,
                },
            },
        )
        raw = _extract_message_text(resp).strip()
        if not raw:
            raise ValueError("Empty model response")
        return json.loads(raw)
    except Exception:
        # Fallback for models/accounts that do not support structured outputs.
        resp = client.chat.completions.create(**base_kwargs)
        raw = _strip_json_fences(_extract_message_text(resp))
        return json.loads(raw)


def analyze_with_ai(full_text: str, api_key: str, style: str = "mla") -> dict:
    """Use OpenAI to detect essay structure and style compliance signals."""
    client = _build_openai_client(api_key)
    if style == "apa":
        prompt = f"""Analyze this essay text and return JSON with these fields:
- "title": essay title (best guess, or "" if unclear)
- "author": author/student name if detectable, else ""
- "institution": institution/affiliation if detectable, else ""
- "instructor": instructor name if detectable, else ""
- "course": course name if detectable, else ""
- "date": due date if detectable, else ""
- "has_references": true/false whether a References/Bibliography section exists
- "title_line_index": 0-based index of the title line (-1 if unsure)
- "heading_end_index": 0-based index of an initial non-body header block (if present), else -1
- "references_start_index": 0-based index where references begin, else -1
- "issues": list of APA 7 student-paper concerns you can infer from text only

Only return valid JSON, nothing else.

Essay text (paragraphs separated by \\n---\\n):
{full_text[:8000]}"""
        schema = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "title": {"type": "string"},
                "author": {"type": "string"},
                "institution": {"type": "string"},
                "instructor": {"type": "string"},
                "course": {"type": "string"},
                "date": {"type": "string"},
                "has_references": {"type": "boolean"},
                "title_line_index": {"type": "integer"},
                "heading_end_index": {"type": "integer"},
                "references_start_index": {"type": "integer"},
                "issues": {"type": "array", "items": {"type": "string"}},
            },
            "required": [
                "title",
                "author",
                "institution",
                "instructor",
                "course",
                "date",
                "has_references",
                "title_line_index",
                "heading_end_index",
                "references_start_index",
                "issues",
            ],
        }
        return _chat_completion_json(client, prompt, "apa_structure_analysis", schema)
    else:
        prompt = f"""Analyze this essay text and return JSON with these fields:
- "title": the essay title (best guess, or "" if unclear)
- "author": author name if detectable, else ""
- "instructor": instructor name if detectable, else ""
- "course": course name if detectable, else ""
- "date": date if detectable, else ""
- "has_works_cited": true/false whether a works cited / bibliography section exists
- "title_line_index": 0-based index of which paragraph is the title (-1 if unsure)
- "heading_end_index": 0-based index of the last heading line (name/instructor/course/date block), -1 if none found
- "works_cited_start_index": 0-based index where works cited begins, -1 if none
- "issues": list of strings describing any MLA compliance problems you notice

Only return valid JSON, nothing else.

Essay text (paragraphs separated by \\n---\\n):
{full_text[:8000]}"""
        schema = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "title": {"type": "string"},
                "author": {"type": "string"},
                "instructor": {"type": "string"},
                "course": {"type": "string"},
                "date": {"type": "string"},
                "has_works_cited": {"type": "boolean"},
                "title_line_index": {"type": "integer"},
                "heading_end_index": {"type": "integer"},
                "works_cited_start_index": {"type": "integer"},
                "issues": {"type": "array", "items": {"type": "string"}},
            },
            "required": [
                "title",
                "author",
                "instructor",
                "course",
                "date",
                "has_works_cited",
                "title_line_index",
                "heading_end_index",
                "works_cited_start_index",
                "issues",
            ],
        }
        return _chat_completion_json(client, prompt, "mla_structure_analysis", schema)


def post_check_content_with_ai(
    body_text: str,
    has_reference_page: bool,
    citation_signals: dict,
    api_key: str,
    style: str = "mla",
) -> list[str]:
    """After formatting, ask AI to review only source/citation compliance (not layout)."""
    client = _build_openai_client(api_key)
    if style == "apa":
        style_label = "APA 7 student-paper"
        citation_label = "APA-style in-text citations"
        ref_label = "References page"
    else:
        style_label = "MLA 9th edition"
        citation_label = "MLA-style parenthetical citations"
        ref_label = "Works Cited page"

    prompt = f"""You are a {style_label} citation/content reviewer.
Analyze only source usage and in-text citation compliance.
Do NOT evaluate or mention margins, spacing, font, page numbers, running headers, title alignment, or paper layout.
Return a JSON array of short warning strings. If no citation/source concerns remain, return [].

Known context:
- {ref_label} detected: {has_reference_page}
- Detected {citation_label}: {citation_signals.get("parenthetical_citation_count", 0)}
- Source signal counts:
  - URLs: {citation_signals.get("url_count", 0)}
  - DOIs: {citation_signals.get("doi_count", 0)}
  - Source cue phrases: {citation_signals.get("source_phrase_count", 0)}
  - Attributed quotes: {citation_signals.get("attributed_quote_count", 0)}
  - Numeric-claim cues: {citation_signals.get("numeric_claim_count", 0)}

Essay body text (first 6000 chars):
{body_text[:6000]}"""
    schema = {
        "type": "array",
        "items": {"type": "string"},
    }
    result = _chat_completion_json(client, prompt, f"{style}_content_warnings", schema)
    return result if isinstance(result, list) else []


def rewrite_works_cited_with_ai(entries: list[str], api_key: str) -> list[str]:
    """(#3) Use gpt-5.2 to rewrite each Works Cited entry to proper MLA 9th ed format."""
    client = _build_openai_client(api_key)

    numbered = "\n".join(f"{i+1}. {e}" for i, e in enumerate(entries))
    prompt = f"""You are an MLA 9th edition citation expert. Below are Works Cited entries that may have formatting errors.
Rewrite each entry to be perfectly MLA 9th edition compliant (correct author order, italics markers, punctuation, hanging indent text, URLs, DOIs, etc.).
Return a JSON array of strings — one corrected entry per element, in the same order. Do not add or remove entries.
Use *asterisks* around text that should be italicized (book/journal titles).

Entries:
{numbered}"""
    schema = {
        "type": "array",
        "items": {"type": "string"},
    }
    result = _chat_completion_json(client, prompt, "mla_works_cited_rewrite", schema)
    # Safety: return original if AI returned wrong count
    return result if isinstance(result, list) and len(result) == len(entries) else entries


# ── Formatting Helpers ───────────────────────────────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
INCH_TOLERANCE = 0.01
PT_TOLERANCE = 0.1
LINE_SPACING_TOLERANCE = 0.05
VALID_STYLES = {"mla", "apa"}


def normalize_style(style: str) -> str:
    value = (style or "mla").strip().lower()
    return value if value in VALID_STYLES else "mla"

def set_run_font(run, bold=False, italic=None):
    """Apply MLA font to a run, preserving italic unless overridden."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = bold
    # Preserve original italic if not explicitly set
    if italic is not None:
        run.font.italic = italic
    run.font.underline = False
    # Force font for East Asian / complex script
    rpr = run._element.get_or_add_rPr()
    for tag in [qn("w:rFonts")]:
        el = rpr.find(tag)
        if el is None:
            el = OxmlElement("w:rFonts")
            rpr.append(el)
        el.set(qn("w:ascii"), FONT_NAME)
        el.set(qn("w:hAnsi"), FONT_NAME)
        el.set(qn("w:cs"), FONT_NAME)

def set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0.5),
                         left_indent=Inches(0), right_indent=Inches(0),
                         space_before=Pt(0), space_after=Pt(0), line_spacing=2.0):
    """Apply MLA paragraph formatting."""
    pf = para.paragraph_format
    pf.alignment = align
    pf.first_line_indent = first_indent
    pf.left_indent = left_indent
    pf.right_indent = right_indent
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.line_spacing = line_spacing
    pf.widow_control = True


def extract_last_name(value: str, default: str = "Author") -> str:
    """Extract a best-effort last name token from a heading/name line."""
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9'’.-]*", (value or "").strip())
    return tokens[-1] if tokens else default


def count_parenthetical_citations_mla(text: str) -> int:
    """Best-effort count of MLA-style parenthetical citations."""
    patterns = [
        r"\([A-Z][A-Za-z'’.-]+(?:\s+(?:and|&)\s+[A-Z][A-Za-z'’.-]+)?\s+\d{1,4}(?:-\d{1,4})?\)",
        r"\((?:qtd\.?\s+in\s+)?[A-Z][A-Za-z'’.-]+(?:\s+et\s+al\.?)?\s+\d{1,4}(?:-\d{1,4})?\)",
        r"\([A-Z][A-Za-z'’.-]+\)",
    ]
    spans = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            spans.add((match.start(), match.end()))
    return len(spans)


def count_parenthetical_citations_apa(text: str) -> int:
    """Best-effort count of APA-style in-text citations."""
    patterns = [
        r"\([A-Z][A-Za-z'’.-]+(?:\s*&\s*[A-Z][A-Za-z'’.-]+)?\,\s*\d{4}[a-z]?(?:,\s*p{1,2}\.?\s*\d+(?:-\d+)?)?\)",
        r"\([A-Z][A-Za-z'’.-]+\s+et\s+al\.,\s*\d{4}[a-z]?(?:,\s*p{1,2}\.?\s*\d+(?:-\d+)?)?\)",
        r"\b[A-Z][A-Za-z'’.-]+\s*\(\d{4}[a-z]?\)",
    ]
    spans = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            spans.add((match.start(), match.end()))
    return len(spans)


def detect_source_signals(text: str, style: str = "mla") -> dict:
    """Heuristics for whether essay body likely relies on sources."""
    url_count = len(re.findall(r"\b(?:https?://|www\.)\S+\b", text, flags=re.IGNORECASE))
    doi_count = len(re.findall(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", text, flags=re.IGNORECASE))
    source_phrase_count = len(
        re.findall(
            r"\b(according to|research shows|research suggests|study|studies|report|reports|survey|data from|statistics|findings)\b",
            text,
            flags=re.IGNORECASE,
        )
    )
    attributed_quote_count = len(
        re.findall(
            r"[\"“][^\"”]{8,}[\"”]\s*(?:\(|,?\s*(?:according to|writes|wrote|argues|argued|states|stated))",
            text,
            flags=re.IGNORECASE,
        )
    )
    numeric_claim_count = len(re.findall(r"\b\d+(?:\.\d+)?%|\b\d{4}\b", text))
    style = normalize_style(style)
    if style == "apa":
        parenthetical_citation_count = count_parenthetical_citations_apa(text)
    else:
        parenthetical_citation_count = count_parenthetical_citations_mla(text)

    likely_source_usage = any(
        [
            url_count > 0,
            doi_count > 0,
            attributed_quote_count > 0,
            source_phrase_count >= 2,
            parenthetical_citation_count > 0,
            (numeric_claim_count >= 3 and source_phrase_count >= 1),
        ]
    )

    return {
        "url_count": url_count,
        "doi_count": doi_count,
        "source_phrase_count": source_phrase_count,
        "attributed_quote_count": attributed_quote_count,
        "numeric_claim_count": numeric_claim_count,
        "parenthetical_citation_count": parenthetical_citation_count,
        "likely_source_usage": likely_source_usage,
    }


def _unique_items(items: list[str]) -> list[str]:
    seen = set()
    ordered = []
    for item in items:
        if not item:
            continue
        if item in seen:
            continue
        seen.add(item)
        ordered.append(item)
    return ordered


def _inches(value) -> float:
    if value is None:
        return 0.0
    try:
        return float(value.inches)
    except Exception:
        return 0.0


def _points(value) -> float:
    if value is None:
        return 0.0
    try:
        return float(value.pt)
    except Exception:
        return 0.0


def _line_spacing_value(para):
    line_spacing = para.paragraph_format.line_spacing
    if line_spacing is None:
        return None
    if isinstance(line_spacing, (int, float)):
        return float(line_spacing)
    try:
        # Convert point-based spacing into line multiples at 12pt baseline.
        return float(line_spacing.pt) / 12.0
    except Exception:
        return None


def _has_page_field(para) -> bool:
    instr_nodes = para._element.findall(".//" + qn("w:instrText"))
    for node in instr_nodes:
        if "PAGE" in (node.text or "").upper():
            return True
    return False


def paragraph_role(idx: int, heading_end: int, title_idx: int, wc_start: int) -> str:
    if heading_end >= 0 and 0 <= idx <= heading_end:
        return "heading"
    if title_idx >= 0 and idx == title_idx:
        return "title"
    if wc_start >= 0 and idx == wc_start:
        return "wc_header"
    if wc_start >= 0 and idx > wc_start:
        return "wc_entry"
    return "body"


def find_works_cited_index(doc) -> int:
    """Deterministically find a Works Cited-style header when AI misses it."""
    headers = {
        "works cited",
        "work cited",
        "references",
        "bibliography",
    }
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip().lower()
        if text in headers:
            return i
    return -1


def is_works_cited_header_text(text: str) -> bool:
    value = (text or "").strip().lower()
    return value in {"works cited", "work cited", "references", "bibliography"}


def is_reference_header_text(text: str) -> bool:
    value = (text or "").strip().lower()
    return value in {"references", "reference", "works cited", "work cited", "bibliography"}


def find_references_index(doc) -> int:
    """Find APA-style references header, with fallback to equivalent bibliography labels."""
    for i, para in enumerate(doc.paragraphs):
        if is_reference_header_text(para.text):
            return i
    return -1


def looks_like_title(text: str) -> bool:
    """Simple heuristic for detecting a standalone title line."""
    value = (text or "").strip()
    if not value:
        return False
    words = value.split()
    if len(words) == 0 or len(words) > 18:
        return False
    if len(value) > 120:
        return False
    if value.endswith((".", "?", "!", ";")):
        return False
    return True


def looks_like_mla_date_line(text: str) -> bool:
    value = (text or "").strip()
    if not value:
        return False
    month_names = (
        r"(January|Jan\.?|February|Feb\.?|March|Mar\.?|April|Apr\.?|May|June|Jun\.?|July|Jul\.?|"
        r"August|Aug\.?|September|Sept?\.?|October|Oct\.?|November|Nov\.?|December|Dec\.?)"
    )
    patterns = [
        rf"^\d{{1,2}}\s+{month_names}\s+\d{{4}}$",
        rf"^{month_names}\s+\d{{1,2}},?\s+\d{{4}}$",
        r"^\d{1,2}/\d{1,2}/\d{2,4}$",
        r"^\d{4}-\d{2}-\d{2}$",
    ]
    return any(re.match(p, value, flags=re.IGNORECASE) for p in patterns)


def guess_mla_heading_end_index(doc) -> int:
    """Heuristic fallback: first 4 non-empty lines where line 4 is date-like."""
    nonempty = [(i, (p.text or "").strip()) for i, p in enumerate(doc.paragraphs) if (p.text or "").strip()]
    if len(nonempty) >= 4 and looks_like_mla_date_line(nonempty[3][1]):
        return nonempty[3][0]
    return -1


def run_verified_mla_checks(doc, last_name: str, heading_end: int, title_idx: int, wc_start: int) -> dict:
    """Deterministic MLA checks that can be verified directly from DOCX properties."""
    passes = []
    warnings = []

    # Margins
    margin_failures = []
    for sec_idx, section in enumerate(doc.sections, start=1):
        for margin_name, value in [
            ("top", section.top_margin),
            ("bottom", section.bottom_margin),
            ("left", section.left_margin),
            ("right", section.right_margin),
        ]:
            if abs(_inches(value) - 1.0) > INCH_TOLERANCE:
                margin_failures.append(f"section {sec_idx} {margin_name}={_inches(value):.2f}\"")
    if margin_failures:
        warnings.append(
            "MLA requires 1-inch margins on all sides. Offending sections: " + ", ".join(margin_failures[:6])
        )
    else:
        passes.append("Verified: 1-inch margins are set on all sections.")

    # Running header
    header_failures = []
    for sec_idx, section in enumerate(doc.sections, start=1):
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else None
        if p is None:
            header_failures.append(f"section {sec_idx}: missing header paragraph")
            continue
        if p.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            header_failures.append(f"section {sec_idx}: header not right-aligned")
        if last_name and last_name.lower() not in (p.text or "").lower():
            header_failures.append(f"section {sec_idx}: header missing last name '{last_name}'")
        if not _has_page_field(p):
            header_failures.append(f"section {sec_idx}: header missing PAGE field")
    if header_failures:
        warnings.append(
            "MLA requires a running header with last name and page number in the upper-right. "
            + "; ".join(header_failures[:6])
        )
    else:
        passes.append("Verified: running header includes last name and page-number field in each section.")

    spacing_before_after_bad = []
    line_spacing_bad = []
    role_format_bad = []
    font_bad = []

    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        role = paragraph_role(idx, heading_end, title_idx, wc_start)
        pf = para.paragraph_format

        if abs(_points(pf.space_before)) > PT_TOLERANCE or abs(_points(pf.space_after)) > PT_TOLERANCE:
            spacing_before_after_bad.append(idx + 1)

        line_spacing = _line_spacing_value(para)
        if line_spacing is None or abs(line_spacing - 2.0) > LINE_SPACING_TOLERANCE:
            line_spacing_bad.append(idx + 1)

        align = pf.alignment
        first = _inches(pf.first_line_indent)
        left = _inches(pf.left_indent)
        right = _inches(pf.right_indent)

        if role == "heading":
            if align != WD_ALIGN_PARAGRAPH.LEFT or abs(first) > INCH_TOLERANCE:
                role_format_bad.append(idx + 1)
        elif role == "title":
            if (
                align != WD_ALIGN_PARAGRAPH.CENTER
                or abs(first) > INCH_TOLERANCE
                or abs(left) > INCH_TOLERANCE
                or abs(right) > INCH_TOLERANCE
            ):
                role_format_bad.append(idx + 1)
        elif role == "wc_header":
            if align != WD_ALIGN_PARAGRAPH.CENTER or abs(first) > INCH_TOLERANCE:
                role_format_bad.append(idx + 1)
        elif role == "wc_entry":
            if align != WD_ALIGN_PARAGRAPH.LEFT or abs(first - (-0.5)) > INCH_TOLERANCE or abs(left - 0.5) > INCH_TOLERANCE:
                role_format_bad.append(idx + 1)
        else:
            if align != WD_ALIGN_PARAGRAPH.LEFT or abs(first - 0.5) > INCH_TOLERANCE:
                role_format_bad.append(idx + 1)

        for run in para.runs:
            if not run.text:
                continue
            run_name = (run.font.name or "").strip().lower()
            run_size = run.font.size.pt if run.font.size is not None else None
            if run_name != FONT_NAME.lower() or run_size is None or abs(float(run_size) - 12.0) > 0.1:
                font_bad.append(idx + 1)
                break

    if spacing_before_after_bad:
        warnings.append(
            "MLA requires no extra spacing between paragraphs (space before/after = 0). Paragraph lines: "
            + ", ".join(str(i) for i in sorted(set(spacing_before_after_bad))[:8])
        )
    else:
        passes.append("Verified: paragraphs use zero space-before and space-after.")

    if line_spacing_bad:
        warnings.append(
            "MLA requires double spacing throughout. Paragraph lines with non-double spacing: "
            + ", ".join(str(i) for i in sorted(set(line_spacing_bad))[:8])
        )
    else:
        passes.append("Verified: paragraphs are double-spaced.")

    if role_format_bad:
        warnings.append(
            "MLA paragraph role formatting mismatch detected (heading/title/body/Works Cited roles). Paragraph lines: "
            + ", ".join(str(i) for i in sorted(set(role_format_bad))[:8])
        )
    else:
        passes.append("Verified: heading, title, body, and Works Cited paragraph indents/alignment match MLA rules.")

    if font_bad:
        warnings.append(
            "MLA requires a readable 12-pt font (configured as Times New Roman 12 pt). Paragraph lines with font mismatches: "
            + ", ".join(str(i) for i in sorted(set(font_bad))[:8])
        )
    else:
        passes.append("Verified: paragraph text runs are normalized to Times New Roman, 12 pt.")

    return {"passes": passes, "warnings": warnings}


def run_verified_apa_checks(doc, title_page_end: int, body_title_idx: int, refs_start: int) -> dict:
    """Deterministic APA checks that can be verified directly from DOCX properties."""
    passes = []
    warnings = []

    # Margins
    margin_failures = []
    for sec_idx, section in enumerate(doc.sections, start=1):
        for margin_name, value in [
            ("top", section.top_margin),
            ("bottom", section.bottom_margin),
            ("left", section.left_margin),
            ("right", section.right_margin),
        ]:
            if abs(_inches(value) - 1.0) > INCH_TOLERANCE:
                margin_failures.append(f"section {sec_idx} {margin_name}={_inches(value):.2f}\"")
    if margin_failures:
        warnings.append(
            "APA requires 1-inch margins on all sides. Offending sections: " + ", ".join(margin_failures[:6])
        )
    else:
        passes.append("Verified: 1-inch margins are set on all sections.")

    # Header: page number only, right aligned
    header_failures = []
    for sec_idx, section in enumerate(doc.sections, start=1):
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else None
        if p is None:
            header_failures.append(f"section {sec_idx}: missing header paragraph")
            continue
        if p.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            header_failures.append(f"section {sec_idx}: header not right-aligned")
        if not _has_page_field(p):
            header_failures.append(f"section {sec_idx}: header missing PAGE field")
    if header_failures:
        warnings.append(
            "APA requires page numbers in the upper-right header. " + "; ".join(header_failures[:6])
        )
    else:
        passes.append("Verified: page-number header exists in each section.")

    spacing_before_after_bad = []
    line_spacing_bad = []
    role_format_bad = []
    font_bad = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        pf = para.paragraph_format

        if abs(_points(pf.space_before)) > PT_TOLERANCE or abs(_points(pf.space_after)) > PT_TOLERANCE:
            spacing_before_after_bad.append(idx + 1)

        line_spacing = _line_spacing_value(para)
        if line_spacing is None or abs(line_spacing - 2.0) > LINE_SPACING_TOLERANCE:
            line_spacing_bad.append(idx + 1)

        align = pf.alignment
        first = _inches(pf.first_line_indent)
        left = _inches(pf.left_indent)
        right = _inches(pf.right_indent)

        is_title_page = 0 <= idx <= title_page_end
        is_body_title = body_title_idx >= 0 and idx == body_title_idx
        is_ref_header = refs_start >= 0 and idx == refs_start
        is_ref_entry = refs_start >= 0 and idx > refs_start

        if is_title_page or is_body_title:
            if (
                align != WD_ALIGN_PARAGRAPH.CENTER
                or abs(first) > INCH_TOLERANCE
                or abs(left) > INCH_TOLERANCE
                or abs(right) > INCH_TOLERANCE
            ):
                role_format_bad.append(idx + 1)
        elif is_ref_header:
            if (
                align != WD_ALIGN_PARAGRAPH.CENTER
                or abs(first) > INCH_TOLERANCE
                or abs(left) > INCH_TOLERANCE
                or abs(right) > INCH_TOLERANCE
            ):
                role_format_bad.append(idx + 1)
        elif is_ref_entry:
            if (
                align != WD_ALIGN_PARAGRAPH.LEFT
                or abs(first - (-0.5)) > INCH_TOLERANCE
                or abs(left - 0.5) > INCH_TOLERANCE
            ):
                role_format_bad.append(idx + 1)
        else:
            if align != WD_ALIGN_PARAGRAPH.LEFT or abs(first - 0.5) > INCH_TOLERANCE:
                role_format_bad.append(idx + 1)

        for run in para.runs:
            if not run.text:
                continue
            run_name = (run.font.name or "").strip().lower()
            run_size = run.font.size.pt if run.font.size is not None else None
            if run_name != FONT_NAME.lower() or run_size is None or abs(float(run_size) - 12.0) > 0.1:
                font_bad.append(idx + 1)
                break

    if spacing_before_after_bad:
        warnings.append(
            "APA requires no extra spacing between paragraphs (space before/after = 0). Paragraph lines: "
            + ", ".join(str(i) for i in sorted(set(spacing_before_after_bad))[:8])
        )
    else:
        passes.append("Verified: paragraphs use zero space-before and space-after.")

    if line_spacing_bad:
        warnings.append(
            "APA requires double spacing throughout. Paragraph lines with non-double spacing: "
            + ", ".join(str(i) for i in sorted(set(line_spacing_bad))[:8])
        )
    else:
        passes.append("Verified: paragraphs are double-spaced.")

    if role_format_bad:
        warnings.append(
            "APA paragraph role formatting mismatch detected (title-page/body/references roles). Paragraph lines: "
            + ", ".join(str(i) for i in sorted(set(role_format_bad))[:8])
        )
    else:
        passes.append("Verified: title page, body, and references paragraph alignment/indents match APA rules.")

    if font_bad:
        warnings.append(
            "APA requires a readable 12-pt font (configured as Times New Roman 12 pt). Paragraph lines with font mismatches: "
            + ", ".join(str(i) for i in sorted(set(font_bad))[:8])
        )
    else:
        passes.append("Verified: paragraph text runs are normalized to Times New Roman, 12 pt.")

    return {"passes": passes, "warnings": warnings}


def set_margins(doc):
    """Set 1-inch margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def set_normal_style_defaults(doc):
    """Normalize Normal style so Word templates can't override MLA spacing/font defaults."""
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        return

    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.line_spacing = 2.0

def add_header(doc, last_name: str):
    """Add right-aligned running header: LastName PageNumber (MLA 9th ed)."""
    for section in doc.sections:
        # Header 0.5" from top edge per MLA
        section.header_distance = Inches(0.5)
        header = section.header
        header.is_linked_to_previous = False
        # Clear existing header content
        for p in header.paragraphs:
            p.clear()
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # MLA: double-spaced, no extra spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(f"{last_name} ")
        set_run_font(run)
        # Page number field
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run2 = p.add_run()
        set_run_font(run2)
        run2._element.append(fld_char1)
        run3 = p.add_run()
        run3._element.append(instr)
        run4 = p.add_run()
        run4._element.append(fld_char2)


def add_page_number_header(doc):
    """Add APA-style page-number-only header (upper-right)."""
    for section in doc.sections:
        section.header_distance = Inches(0.5)
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run()
        set_run_font(run)
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._element.append(fld_char1)
        run2 = p.add_run()
        run2._element.append(instr)
        run3 = p.add_run()
        run3._element.append(fld_char2)


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)
    set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0))


def add_text_paragraph(doc, text: str, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0.5), bold=False,
                       left_indent=Inches(0), right_indent=Inches(0)):
    para = doc.add_paragraph(text or "")
    set_paragraph_format(
        para,
        align=align,
        first_indent=first_indent,
        left_indent=left_indent,
        right_indent=right_indent,
    )
    if not para.runs:
        para.add_run("")
    for run in para.runs:
        set_run_font(run, bold=bold)
    return para


def add_heading_block(doc, name, instructor, course, date, heading_order=None):
    """Insert MLA first-page heading block. Default: Name / Instructor / Course / Date.
    heading_order overrides the line order (list of field keys)."""
    ref_para = doc.paragraphs[0] if doc.paragraphs else None
    field_map = {"name": name, "instructor": instructor, "course": course, "date": date}
    # (#8) Use custom order if provided, otherwise default MLA order
    if heading_order and len(heading_order) == 4:
        lines = [field_map.get(k, "") for k in heading_order]
    else:
        lines = [name, instructor, course, date]
    for line in lines:
        new_p = OxmlElement("w:p")
        # Paragraph properties: left-aligned, double-spaced, no indent, no extra spacing
        pPr = OxmlElement("w:pPr")
        # Normalize style to Normal
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Normal")
        pPr.append(pStyle)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "480")      # 480 twips = double-spaced 12pt
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), "0")
        pPr.append(ind)
        new_p.append(pPr)
        # Run with text + font properties
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)
        rFonts.set(qn("w:cs"), FONT_NAME)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")  # 24 half-points = 12pt
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "24")
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)
        new_p.append(r)
        if ref_para is not None:
            ref_para._element.addprevious(new_p)
        else:
            doc.element.body.append(new_p)


def remove_empty_paragraphs(doc, heading_end, title_idx, wc_start):
    """(#1) Remove empty paragraphs — MLA has no extra blank lines."""
    to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            continue
        to_remove.append(i)
    # Remove in reverse to keep indices stable
    for i in reversed(to_remove):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    def remap_index(idx: int) -> int:
        if idx < 0 or idx in to_remove:
            return -1
        return idx - sum(1 for r in to_remove if r < idx)

    # If heading_end itself was removed, move back to previous surviving heading line.
    if heading_end >= 0:
        while heading_end >= 0 and heading_end in to_remove:
            heading_end -= 1

    new_title = remap_index(title_idx)
    new_wc = remap_index(wc_start)
    new_heading_end = remap_index(heading_end)
    return new_heading_end, new_title, new_wc

def add_page_break_before(para):
    """Insert a page break before a paragraph."""
    pPr = para._element.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    page_break.set(qn("w:val"), "true")
    pPr.append(page_break)


# ── Core Formatter ───────────────────────────────────────────────────────────


def get_body_text(doc, heading_end: int, title_idx: int, wc_start: int) -> str:
    lines = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if paragraph_role(idx, heading_end, title_idx, wc_start) in {"heading", "title", "wc_header", "wc_entry"}:
            continue
        lines.append(text)
    return "\n".join(lines)


def format_apa(
    input_path: str,
    output_path: str,
    name="",
    instructor="",
    course="",
    date="",
    institution="",
    use_ai=True,
    api_key=None,
):
    """Format document to APA 7 student-paper defaults."""
    source_doc = Document(input_path)
    para_texts = [p.text.strip() for p in source_doc.paragraphs]
    full_text = "\n---\n".join(para_texts)

    result = {
        "title": "",
        "pre_issues": [],
        "post_issues": [],
        "preview": [],
        "verified_passes": [],
        "verified_warnings": [],
        "content_warnings": [],
        "manual_review": [],
    }

    name = (name or "").strip()
    instructor = (instructor or "").strip()
    course = (course or "").strip()
    date = (date or "").strip()
    institution = (institution or "").strip()

    ai_info = None
    if use_ai and api_key:
        try:
            ai_info = analyze_with_ai(full_text, api_key, style="apa")
            result["title"] = (ai_info.get("title") or "").strip()
            result["pre_issues"] = ai_info.get("issues", [])
            name = name or (ai_info.get("author") or "").strip()
            institution = institution or (ai_info.get("institution") or "").strip()
            instructor = instructor or (ai_info.get("instructor") or "").strip()
            course = course or (ai_info.get("course") or "").strip()
            date = date or (ai_info.get("date") or "").strip()
        except Exception:
            pass

    heading_end = ai_info.get("heading_end_index", -1) if ai_info else -1
    title_idx = ai_info.get("title_line_index", -1) if ai_info else -1
    refs_start = ai_info.get("references_start_index", -1) if ai_info else -1
    if refs_start < 0:
        refs_start = find_references_index(source_doc)

    # Fallback title detection for non-AI and imperfect AI cases.
    if title_idx < 0 or title_idx >= len(source_doc.paragraphs):
        search_start = 0
        if heading_end >= 0:
            search_start = min(heading_end + 1, len(source_doc.paragraphs))
        elif len(source_doc.paragraphs) >= 4 and looks_like_mla_date_line(source_doc.paragraphs[3].text):
            search_start = 4
        for i in range(search_start, len(source_doc.paragraphs)):
            text = source_doc.paragraphs[i].text.strip()
            if text and not is_reference_header_text(text) and looks_like_title(text):
                title_idx = i
                break

    if not result["title"] and 0 <= title_idx < len(source_doc.paragraphs):
        result["title"] = source_doc.paragraphs[title_idx].text.strip()

    first_nonempty_idx = next((i for i, t in enumerate(para_texts) if t), -1)
    if not result["title"] and first_nonempty_idx >= 0 and looks_like_title(para_texts[first_nonempty_idx]):
        result["title"] = para_texts[first_nonempty_idx]
        title_idx = first_nonempty_idx

    title = (result["title"] or "Untitled Essay").strip()
    result["title"] = title

    body_start = 0
    if title_idx >= 0:
        body_start = title_idx + 1
    elif heading_end >= 0:
        body_start = heading_end + 1
    elif len(source_doc.paragraphs) >= 5 and looks_like_mla_date_line(source_doc.paragraphs[3].text):
        # Common MLA input shape: 4 heading lines + title + body.
        body_start = 5

    body_end = refs_start if refs_start >= 0 else len(source_doc.paragraphs)
    body_lines = []
    skipped_duplicate_title = False
    for i in range(max(0, body_start), max(0, body_end)):
        text = para_texts[i].strip()
        if not text:
            continue
        if not skipped_duplicate_title and text.lower() == title.lower():
            skipped_duplicate_title = True
            continue
        body_lines.append(text)

    reference_entries = []
    if 0 <= refs_start < len(source_doc.paragraphs):
        for i in range(refs_start + 1, len(source_doc.paragraphs)):
            text = para_texts[i].strip()
            if text:
                reference_entries.append(text)

    # Build clean APA document.
    doc = Document()
    set_normal_style_defaults(doc)
    set_margins(doc)
    add_page_number_header(doc)

    # Title page content starts around upper half by adding leading blank lines.
    for _ in range(3):
        add_text_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0))
    add_text_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0), bold=True)
    if any([name, institution, course, instructor, date]):
        add_text_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0))
    for line in [name, institution, course, instructor, date]:
        if line:
            add_text_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0))

    title_page_end = len(doc.paragraphs) - 1

    add_page_break(doc)
    add_text_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0), bold=True)
    body_title_idx = len(doc.paragraphs) - 1

    for line in body_lines:
        add_text_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0.5))

    refs_start_new = -1
    if reference_entries:
        add_page_break(doc)
        add_text_paragraph(doc, "References", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0), bold=True)
        refs_start_new = len(doc.paragraphs) - 1
        for entry in reference_entries:
            add_text_paragraph(
                doc,
                entry,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_indent=Inches(-0.5),
                left_indent=Inches(0.5),
            )

    # Preview payload
    preview = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if 0 <= i <= title_page_end:
            ptype = "apa_titlepage"
        elif i == body_title_idx:
            ptype = "apa_body_title"
        elif refs_start_new >= 0 and i == refs_start_new:
            ptype = "ref_header"
        elif refs_start_new >= 0 and i > refs_start_new:
            ptype = "ref_entry"
        else:
            ptype = "body"
        preview.append(
            {
                "type": ptype,
                "text": text,
                "pb": (i == body_title_idx) or (refs_start_new >= 0 and i == refs_start_new),
            }
        )
    result["preview"] = preview
    result["last_name"] = ""

    verified = run_verified_apa_checks(doc, title_page_end, body_title_idx, refs_start_new)
    result["verified_passes"] = verified["passes"]
    result["verified_warnings"] = verified["warnings"]

    has_references = refs_start_new >= 0 and len(reference_entries) > 0
    body_text = "\n".join(body_lines)
    citation_signals = detect_source_signals(body_text, style="apa")
    ai_content_warnings = []
    ai_content_check_failed = False
    if use_ai and api_key:
        try:
            ai_content_warnings = post_check_content_with_ai(
                body_text,
                has_references,
                citation_signals,
                api_key,
                style="apa",
            )
        except Exception:
            ai_content_check_failed = True

    content_warnings = []
    if not has_references and citation_signals.get("likely_source_usage", False):
        content_warnings.append("No References page detected, but source usage appears likely.")
    if has_references and citation_signals.get("likely_source_usage", False) and citation_signals.get("parenthetical_citation_count", 0) == 0:
        content_warnings.append("Potential missing APA in-text citations in body text.")
    content_warnings.extend(ai_content_warnings)
    result["content_warnings"] = _unique_items(content_warnings)

    manual_review = []
    if not has_references and not citation_signals.get("likely_source_usage", False):
        manual_review.append("No References detected; acceptable only if assignment permits no-source essays.")
    if ai_content_check_failed:
        manual_review.append("AI citation review was unavailable; manually confirm source usage and in-text citation compliance.")
    result["manual_review"] = _unique_items(manual_review)

    result["post_issues"] = _unique_items(
        result["verified_warnings"] + result["content_warnings"] + result["manual_review"]
    )

    doc.save(output_path)
    return result


def format_document(input_path: str, output_path: str, style: str = "mla", **kwargs):
    style = normalize_style(style)
    if style == "apa":
        kwargs.pop("no_heading", None)
        kwargs.pop("heading_order", None)
        return format_apa(input_path, output_path, **kwargs)
    kwargs.pop("institution", None)
    return format_mla(input_path, output_path, **kwargs)


def format_mla(input_path: str, output_path: str, name="", instructor="", course="", date="",
               use_ai=True, api_key=None, no_heading=False, heading_order=None):
    """Main formatting pipeline. Returns structured result dict."""
    doc = Document(input_path)

    # Keep user-provided heading input distinct from AI fallbacks.
    name = (name or "").strip()
    instructor = (instructor or "").strip()
    course = (course or "").strip()
    date = (date or "").strip()

    # Extract all paragraph texts for AI analysis
    para_texts = [p.text.strip() for p in doc.paragraphs]
    full_text = "\n---\n".join(para_texts)

    result = {
        "title": "",
        "pre_issues": [],
        "post_issues": [],
        "preview": [],
        "verified_passes": [],
        "verified_warnings": [],
        "content_warnings": [],
        "manual_review": [],
    }
    ai_info = None

    # Track user intent (blank form vs provided heading fields) before AI fallback.
    _user_name = name
    _user_heading_supplied = any([name, instructor, course, date])

    if use_ai and api_key:
        try:
            ai_info = analyze_with_ai(full_text, api_key, style="mla")
            result["title"] = ai_info.get("title", "")
            result["pre_issues"] = ai_info.get("issues", [])
            # Use AI-detected info as fallback for heading fields
            name = name or ai_info.get("author", "")
            instructor = instructor or ai_info.get("instructor", "")
            course = course or ai_info.get("course", "")
            date = date or ai_info.get("date", "")
        except Exception:
            pass

    # Detect structure indices from AI
    ai_heading_end = ai_info.get("heading_end_index", -1) if ai_info else -1
    title_idx = ai_info.get("title_line_index", -1) if ai_info else -1
    wc_start = ai_info.get("works_cited_start_index", -1) if ai_info else -1
    if ai_heading_end < 0:
        ai_heading_end = guess_mla_heading_end_index(doc)
    if wc_start < 0:
        wc_start = find_works_cited_index(doc)

    # Derive last_name for running header
    if _user_name:
        last_name = extract_last_name(_user_name)
    elif not _user_heading_supplied and ai_heading_end >= 0 and doc.paragraphs:
        first_heading_line = ""
        for i in range(min(ai_heading_end + 1, len(doc.paragraphs))):
            candidate = doc.paragraphs[i].text.strip()
            if candidate:
                first_heading_line = candidate
                break
        last_name = extract_last_name(first_heading_line)
    elif name:
        last_name = extract_last_name(name)
    else:
        last_name = "Author"

    # ── Apply formatting ──
    set_normal_style_defaults(doc)
    set_margins(doc)
    add_header(doc, last_name)

    if no_heading:
        heading_end = -1
    elif _user_heading_supplied:
        # User explicitly provided info: remove old heading, insert fresh MLA heading block
        if ai_heading_end >= 0:
            for idx in range(ai_heading_end, -1, -1):
                doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
            removed = ai_heading_end + 1
            title_idx = max(title_idx - removed, -1) if title_idx >= 0 else -1
            wc_start = max(wc_start - removed, -1) if wc_start >= 0 else -1
        add_heading_block(doc, name, instructor, course, date, heading_order=heading_order)
        heading_end = 3  # always 4 lines (indices 0-3)
        if title_idx >= 0:
            title_idx += 4
        else:
            title_idx = 4
        if wc_start >= 0:
            wc_start += 4
    else:
        # No user input: keep existing heading, just reformat it in place
        heading_end = ai_heading_end

    # (#1) Remove empty paragraphs (MLA has no extra blank lines)
    heading_end, title_idx, wc_start = remove_empty_paragraphs(doc, heading_end, title_idx, wc_start)

    # AI can point wc_start to an empty spacer paragraph. Re-detect on cleaned doc when needed.
    if wc_start < 0 or wc_start >= len(doc.paragraphs) or not is_works_cited_header_text(doc.paragraphs[wc_start].text):
        wc_start = find_works_cited_index(doc)

    # Title is always right after heading — override AI's unreliable index
    if heading_end >= 0 and heading_end + 1 < len(doc.paragraphs):
        title_idx = heading_end + 1
    elif title_idx < 0 or title_idx >= len(doc.paragraphs):
        # Fallback for docs without detected heading: promote first plausible short line to title.
        search_start = 0
        if len(doc.paragraphs) >= 4 and looks_like_mla_date_line(doc.paragraphs[3].text):
            # AI may have missed a 4-line MLA heading block; skip it before finding title.
            search_start = 4
        for i in range(search_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if looks_like_title(para.text):
                title_idx = i
                break

    # Page break before Works Cited
    if wc_start >= 0 and wc_start < len(doc.paragraphs):
        add_page_break_before(doc.paragraphs[wc_start])

    # Format each paragraph
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Normalize paragraph style to Normal
        try:
            para.style = doc.styles['Normal']
        except KeyError:
            pass

        role = paragraph_role(i, heading_end, title_idx, wc_start)
        is_title = role == "title"

        if role == "heading":
            # Left-aligned, no indent
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0))
        elif role == "title":
            # Centered, no indent, not bold, not italic
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0))
        elif role == "wc_header":
            # "Works Cited" centered, no indent
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Inches(0))
        elif role == "wc_entry":
            # Hanging indent: -0.5" first-line, 0.5" left indent
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(-0.5))
            para.paragraph_format.left_indent = Inches(0.5)
        else:
            # Body paragraph: 0.5" first-line indent
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0.5))

        # Apply font to all runs
        for run in para.runs:
            if is_title:
                # Title: no bold, no italic, no underline
                set_run_font(run, bold=False, italic=False)
            else:
                set_run_font(run, bold=False)

    # (#3) Rewrite Works Cited entries with AI (gpt-5.2) for proper MLA formatting
    if use_ai and api_key and wc_start >= 0:
        wc_entries = []
        wc_indices = []
        for i, para in enumerate(doc.paragraphs):
            if wc_start >= 0 and i > wc_start and para.text.strip():
                wc_entries.append(para.text.strip())
                wc_indices.append(i)
        if wc_entries:
            try:
                corrected = rewrite_works_cited_with_ai(wc_entries, api_key)
                for idx, new_text in zip(wc_indices, corrected):
                    para = doc.paragraphs[idx]
                    # Clear existing runs and replace with corrected text
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        # Handle *italics* markers from AI
                        parts = re.split(r'(\*[^*]+\*)', new_text)
                        para.runs[0].text = ""
                        first = True
                        for part in parts:
                            if not part:
                                continue
                            if part.startswith('*') and part.endswith('*'):
                                r = para.add_run(part[1:-1])
                                set_run_font(r, bold=False, italic=True)
                            else:
                                if first and para.runs:
                                    para.runs[0].text = part
                                    first = False
                                else:
                                    r = para.add_run(part)
                                    set_run_font(r, bold=False)
                    else:
                        para.add_run(new_text)
                        set_run_font(para.runs[0], bold=False)
            except Exception:
                pass  # Keep original entries on failure

    # Build structured preview — ALL paragraphs for full-doc page rendering
    result["last_name"] = last_name
    preview = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        role = paragraph_role(i, heading_end, title_idx, wc_start)
        # Mark page break before Works Cited
        page_break = (i == wc_start) if wc_start >= 0 else False
        if role == "heading":
            ptype = "heading"
        elif role == "title":
            ptype = "title"
        elif role == "wc_header":
            ptype = "wc_header"
        elif role == "wc_entry":
            ptype = "wc_entry"
        else:
            ptype = "body"
        preview.append({"type": ptype, "text": t, "pb": page_break})
    result["preview"] = preview

    # Deterministic, directly verifiable MLA checks.
    verified = run_verified_mla_checks(doc, last_name, heading_end, title_idx, wc_start)
    result["verified_passes"] = verified["passes"]
    result["verified_warnings"] = verified["warnings"]

    # Source/citation heuristics + AI content review (layout excluded).
    has_works_cited = wc_start >= 0
    body_text = get_body_text(doc, heading_end, title_idx, wc_start)
    citation_signals = detect_source_signals(body_text, style="mla")
    ai_content_warnings = []
    ai_content_check_failed = False
    if use_ai and api_key:
        try:
            ai_content_warnings = post_check_content_with_ai(
                body_text,
                has_works_cited,
                citation_signals,
                api_key,
                style="mla",
            )
        except Exception:
            ai_content_check_failed = True

    content_warnings = []
    if not has_works_cited and citation_signals.get("likely_source_usage", False):
        content_warnings.append("No Works Cited page detected, but source usage appears likely.")
    if has_works_cited and citation_signals.get("likely_source_usage", False) and citation_signals.get("parenthetical_citation_count", 0) == 0:
        content_warnings.append("Potential missing MLA parenthetical citations in body text.")
    content_warnings.extend(ai_content_warnings)
    result["content_warnings"] = _unique_items(content_warnings)

    manual_review = []
    if not has_works_cited and not citation_signals.get("likely_source_usage", False):
        manual_review.append("No Works Cited detected; acceptable only if assignment permits no-source essays.")
    if ai_content_check_failed:
        manual_review.append("AI citation review was unavailable; manually confirm source usage and in-text citation compliance.")
    result["manual_review"] = _unique_items(manual_review)

    # Backward-compatible list consumed by existing UI.
    result["post_issues"] = _unique_items(
        result["verified_warnings"] + result["content_warnings"] + result["manual_review"]
    )

    doc.save(output_path)
    return result


# ── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Reformat a DOCX essay to MLA 9 or APA 7")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("-o", "--output", help="Output .docx file (default: input_<style>.docx)")
    parser.add_argument("--style", default="mla", choices=["mla", "apa"], help="Formatting style")
    parser.add_argument("--no-ai", action="store_true", help="Skip AI analysis")
    parser.add_argument("--no-heading", action="store_true", help="Skip heading insertion")
    parser.add_argument("--name", default="", help="Student name")
    parser.add_argument("--institution", default="", help="Institution (APA title page)")
    parser.add_argument("--instructor", default="", help="Instructor name")
    parser.add_argument("--course", default="", help="Course name")
    parser.add_argument("--date", default="", help="Date string")
    parser.add_argument("--api-key", default=None, help="OpenAI API key (or set OPENAI_API_KEY env var)")
    args = parser.parse_args()

    if not args.input.endswith(".docx"):
        sys.exit("Error: input must be a .docx file")

    style = normalize_style(args.style)
    output = args.output or args.input.replace(".docx", f"_{style}.docx")
    api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "")
    use_ai = not args.no_ai and bool(api_key)

    if not args.no_ai and not api_key:
        print("No OpenAI API key found. Set OPENAI_API_KEY or pass --api-key. Running without AI.\n")

    result = format_document(
        args.input,
        output,
        style=style,
        name=args.name,
        institution=args.institution,
        instructor=args.instructor,
        course=args.course,
        date=args.date,
        use_ai=use_ai,
        api_key=api_key,
        no_heading=args.no_heading,
    )

    if result.get("title"):
        print(f"Detected title: \"{result['title']}\"")
    if result.get("pre_issues"):
        print("Pre-format issues:")
        for issue in result["pre_issues"]:
            print(f"  - {issue}")
    if result.get("post_issues") is not None:
        if result["post_issues"]:
            print("Post-format issues:")
            for issue in result["post_issues"]:
                print(f"  - {issue}")
        else:
            print(f"No {style.upper()} issues detected!")
    print(f"\nSaved: {output}")

if __name__ == "__main__":
    main()
